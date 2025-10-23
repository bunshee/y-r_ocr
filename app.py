import io
import os
import zipfile

import pandas as pd
import streamlit as st
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from agent import SelfDescribingOCRAgent

st.title("Young & Restless OCR")


def dataframe_to_pdf(tables_with_pages, display_date=None):
    """Converts a list of (page_num, DataFrame) tuples to a PDF."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()

    # Create a custom cell style with smaller font size
    cell_style = styles["Normal"]
    cell_style.fontSize = 8  # Set font size to 8 (you can adjust this value)
    cell_style.wordWrap = "CJK"  # Enable word wrapping

    # Add titles and date
    elements.append(Paragraph("THE YOUNG AND THE RESTLESS", styles["h1"]))
    elements.append(Paragraph("PRELIMINARY PRODUCTION REPORT", styles["h2"]))
    if display_date:
        elements.append(Paragraph(f"DATE: {display_date}", styles["h3"]))

    for i, (page_num, df) in enumerate(tables_with_pages):
        # Remove empty columns
        df = df.dropna(axis=1, how="all")
        if df.empty:
            continue

        # Convert data to a list of lists of Paragraphs for word wrapping
        # Ensure empty cells are truly empty, not 'None' or 'nan' strings
        processed_data = []
        # Add header row
        processed_data.append(
            [Paragraph(str(col), cell_style) for col in df.columns.to_list()]
        )
        # Add data rows
        for _, row in df.iterrows():
            processed_row = []
            for cell_value in row:
                if pd.isna(cell_value) or str(cell_value).strip() == "":
                    processed_row.append(
                        Paragraph("", cell_style)
                    )  # Empty Paragraph for empty cells
                else:
                    processed_row.append(Paragraph(str(cell_value), cell_style))
            processed_data.append(processed_row)

        # Create table and allow splitting
        num_cols = len(df.columns)

        # Calculate column widths to fit the page
        available_width = (
            letter[0] - doc.leftMargin - doc.rightMargin - 20
        )  # Subtract some margin for safety
        col_widths = []
        if num_cols > 0:
            col_width = available_width / num_cols
            col_widths = [col_width] * num_cols

        table = Table(
            processed_data, colWidths=col_widths if col_widths else None, splitByRow=1
        )

        # Add style
        style = TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]
        )
        table.setStyle(style)

        elements.append(table)
        if i < len(tables_with_pages) - 1:  # Add PageBreak only if not the last table
            elements.append(PageBreak())
        else:
            elements.append(Spacer(1, 0))  # Add a small spacer at the end

    doc.build(elements)
    buffer.seek(0)
    return buffer


def dataframes_to_docx(tables_with_pages, display_date=None):
    """Converts a list of (page_num, DataFrame) tuples to a DOCX file."""
    document = Document()
    document.add_heading("THE YOUNG AND THE RESTLESS", level=1)
    document.add_heading("PRELIMINARY PRODUCTION REPORT", level=2)
    if display_date:
        document.add_heading(f"DATE: {display_date}", level=3)

    for page_num, df in tables_with_pages:
        # Add table
        table = document.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"

        # Add header row
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = str(col_name)

        # Add data rows
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row):
                if cell_value is None:
                    row_cells[i].text = ""
                else:
                    row_cells[i].text = str(cell_value)

    # Save to a byte buffer
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def dataframes_to_xlsx(tables_with_pages, display_date=None):
    """Converts a list of (page_num, DataFrame) tuples to an XLSX file."""
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        if display_date:
            # Add a summary sheet or a note about the date if needed
            summary_df = pd.DataFrame([{"Report Date": display_date}])
            summary_df.to_excel(writer, sheet_name="Report Summary", index=False)

        for page_num, df in tables_with_pages:
            # Ensure sheet name is valid (max 31 chars, no invalid chars)
            sheet_name = f"Page {page_num}"
            if len(sheet_name) > 31:
                sheet_name = sheet_name[:31]
            df.to_excel(writer, sheet_name=sheet_name, index=False)
    buffer.seek(0)
    return buffer


def create_zip_archive(tables_with_pages, date_to_display):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        # Generate XLSX
        if tables_with_pages:
            xlsx_buffer = dataframes_to_xlsx(tables_with_pages, date_to_display)
            zip_file.writestr("extracted_tables.xlsx", xlsx_buffer.read())

        # Generate PDF
        if tables_with_pages:
            pdf_buffer = dataframe_to_pdf(tables_with_pages, date_to_display)
            zip_file.writestr("extracted_tables.pdf", pdf_buffer.read())

        # Generate DOCX
        if tables_with_pages:
            docx_buffer = dataframes_to_docx(tables_with_pages, date_to_display)
            zip_file.writestr("extracted_tables.docx", docx_buffer.read())

    zip_buffer.seek(0)
    return zip_buffer


# Get API key from user
api_key = st.text_input("Enter your Google API Key:", type="password")

# File uploader
uploaded_file = st.file_uploader("Choose a PDF file", type="pdf")

# Initialize variables outside the try block
tables_with_pages = []
date_to_display = None
processing_successful = False

if uploaded_file is not None and api_key:
    # Save the uploaded file to a temporary location
    temp_pdf_path = os.path.abspath("temp.pdf")
    with open(temp_pdf_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    if st.button("Process Document"):
        try:
            # Initialize the agent
            agent = SelfDescribingOCRAgent(api_key=api_key)

            # Run the full pipeline
            results = agent.process(temp_pdf_path)

            edited_tables_with_pages = []

            for page_num, _, line_items, _, corrected_image_bytes in results:
                if not isinstance(line_items, pd.DataFrame):
                    st.warning(
                        f"line_items is not a DataFrame for page {page_num}. Type: {type(line_items)}"
                    )
                    line_items = (
                        pd.DataFrame()
                    )  # Convert to empty DataFrame to prevent further errors

                # Clean the dataframe before displaying
                if not isinstance(line_items, pd.DataFrame):
                    line_items = pd.DataFrame()

                # Function to clean cell values more leniently
                def clean_cell(value):
                    if pd.isna(value) or value == "" or value is None:
                        return None
                    str_val = str(value).strip()
                    # Only return None for truly empty strings after stripping
                    return str_val if str_val else None

                st.subheader(f"Page {page_num}")

                if corrected_image_bytes:
                    st.image(
                        corrected_image_bytes,
                        caption=f"Rotated Image for Page {page_num}",
                        width="stretch",
                    )
                line_items_cleaned = line_items.map(clean_cell)

                line_items_cleaned = line_items_cleaned.dropna(axis=1, how="all")
                line_items_cleaned = line_items_cleaned.dropna(axis=0, how="all")
                # line_items_cleaned = line_items_cleaned.replace("✔", None)
                # line_items_cleaned = line_items_cleaned.fillna(method="ffill")
                # line_items_cleaned = line_items_cleaned.fillna("")
                if (
                    not line_items_cleaned.empty
                    and not line_items_cleaned.dropna(how="all").empty
                ):
                    try:
                        edited_df = st.data_editor(
                            line_items_cleaned, key=f"editor_{page_num}"
                        )
                    except Exception as e:
                        st.warning(
                            f"Could not display dataframe directly, converting to string: {e}"
                        )
                        edited_df = st.data_editor(
                            line_items_cleaned.astype(str), key=f"editor_{page_num}"
                        )
                    line_items_cleaned = edited_df
                else:
                    st.info("No data to display for this page.")
                    continue
                if not line_items_cleaned.empty:
                    edited_tables_with_pages.append((page_num, line_items_cleaned))

                if date_to_display is None:
                    for col in line_items.columns:
                        if "date" in col.lower():
                            for item in line_items[col]:
                                try:
                                    parsed_date = pd.to_datetime(str(item))
                                    date_to_display = parsed_date.strftime(
                                        "%A, %B %d, %Y"
                                    )
                                    break
                                except (ValueError, TypeError):
                                    continue
                            if date_to_display:
                                break
            processing_successful = True

        except Exception as e:
            st.error(f"An error occurred: {e}")

    # Render download buttons outside the try block, but only if processing was successful
    if processing_successful and edited_tables_with_pages:
        zip_file_buffer = create_zip_archive(edited_tables_with_pages, date_to_display)
        st.download_button(
            label="Download All Results (ZIP)",
            data=zip_file_buffer,
            file_name="extracted_results.zip",
            mime="application/zip",
        )
