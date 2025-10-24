import io
import zipfile

import pandas as pd
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


def dataframe_to_pdf(tables_with_pages, display_date=None):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()

    cell_style = styles["Normal"]
    cell_style.fontSize = 8
    cell_style.wordWrap = "CJK"

    elements.append(Paragraph("THE YOUNG AND THE RESTLESS", styles["h1"]))
    elements.append(Paragraph("PRELIMINARY PRODUCTION REPORT", styles["h2"]))
    if display_date:
        elements.append(Paragraph(f"DATE: {display_date}", styles["h3"]))

    for i, (page_num, df) in enumerate(tables_with_pages):
        df = df.dropna(axis=1, how="all")
        if df.empty:
            continue

        processed_data = [
            [Paragraph(str(col), cell_style) for col in df.columns.to_list()]
        ]
        for _, row in df.iterrows():
            processed_row = [
                Paragraph(
                    ""
                    if pd.isna(cell_value) or str(cell_value).strip() == ""
                    else str(cell_value),
                    cell_style,
                )
                for cell_value in row
            ]
            processed_data.append(processed_row)

        table = Table(processed_data, splitByRow=1)

        style = TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                ("GRID", (0, 0), (-1, -1), 1, colors.black),
            ]
        )
        table.setStyle(style)

        elements.append(table)
        if i < len(tables_with_pages) - 1:
            elements.append(PageBreak())
        else:
            elements.append(Spacer(1, 0))

    doc.build(elements)
    buffer.seek(0)
    return buffer


def dataframes_to_docx(tables_with_pages, display_date=None):
    document = Document()
    document.add_heading("THE YOUNG AND THE RESTLESS", level=1)
    document.add_heading("PRELIMINARY PRODUCTION REPORT", level=2)
    if display_date:
        document.add_heading(f"DATE: {display_date}", level=3)

    for page_num, df in tables_with_pages:
        table = document.add_table(rows=1, cols=len(df.columns))
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = str(col_name)

        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row):
                row_cells[i].text = "" if cell_value is None else str(cell_value)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def dataframes_to_xlsx(tables_with_pages, display_date=None):
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        if display_date:
            summary_df = pd.DataFrame([{"Report Date": display_date}])
            summary_df.to_excel(writer, sheet_name="Report Summary", index=False)

        for page_num, df in tables_with_pages:
            sheet_name = f"Page {page_num}"
            if len(sheet_name) > 31:
                sheet_name = sheet_name[:31]
            df.to_excel(writer, sheet_name=sheet_name, index=False)
    buffer.seek(0)
    return buffer


def create_zip_archive(tables_with_pages, date_to_display):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
        if tables_with_pages:
            xlsx_buffer = dataframes_to_xlsx(tables_with_pages, date_to_display)
            zip_file.writestr("extracted_tables.xlsx", xlsx_buffer.read())

        if tables_with_pages:
            pdf_buffer = dataframe_to_pdf(tables_with_pages, date_to_display)
            zip_file.writestr("extracted_tables.pdf", pdf_buffer.read())

        if tables_with_pages:
            docx_buffer = dataframes_to_docx(tables_with_pages, date_to_display)
            zip_file.writestr("extracted_tables.docx", docx_buffer.read())

    zip_buffer.seek(0)
    return zip_buffer
